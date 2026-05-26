from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT = OUT_DIR / "BaoCao_DuAn_BanVeTauNhaGa.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=(255, 255, 255))
        set_cell_shading(table.rows[0].cells[index], "005D90")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value))
    return table


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_page_number_footer(document):
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "BanVeTauNhaGa - Báo cáo dự án | Trang "
    run = footer.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)


def build_document():
    OUT_DIR.mkdir(exist_ok=True)
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(0, 93, 144)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(0, 93, 144)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("TRƯỜNG/ĐƠN VỊ THỰC HIỆN")
    run.bold = True
    run.font.size = Pt(14)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("BÁO CÁO DỰ ÁN")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 93, 144)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Hệ thống quản lý bán vé tàu nhà ga")
    run.bold = True
    run.font.size = Pt(16)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Dự án: BanVeTauNhaGa").bold = True
    paragraph.add_run("\nNền tảng: Java Swing + SQL Server + Maven")
    paragraph.add_run(f"\nNgày lập báo cáo: {datetime.now().strftime('%d/%m/%Y')}")

    document.add_paragraph("\n\n")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Người thực hiện: ........................................").italic = True
    document.add_page_break()

    document.add_heading("1. Tổng quan dự án", level=1)
    document.add_paragraph(
        "BanVeTauNhaGa là ứng dụng desktop phục vụ nghiệp vụ quản lý bán vé tàu tại nhà ga. "
        "Hệ thống được xây dựng bằng Java Swing, tổ chức theo các lớp giao diện, DAO truy cập dữ liệu, "
        "entity mô hình hóa bảng dữ liệu và cơ sở dữ liệu SQL Server."
    )
    document.add_paragraph(
        "Mục tiêu chính của dự án là hỗ trợ nhân viên nhà ga thao tác bán vé, quản lý vé, hóa đơn, "
        "lịch chạy, tuyến, đoàn tàu, toa, ghế, khách hàng, nhân viên, giá vé và khuyến mãi trong một giao diện thống nhất."
    )
    add_table(document, ["Hạng mục", "Thông tin"], [
        ["Tên artifact", "BanVeTauNhaGa"],
        ["Ngôn ngữ/chạy chính", "Java 17, Swing"],
        ["Cơ sở dữ liệu", "Microsoft SQL Server - BanVeTauNhaGa"],
        ["Cách build", "Maven, main class com.Main"],
        ["Số file Java khảo sát", "113 file"],
        ["Số module UI", "63 file trong com.modules"],
        ["Số DAO", "21 file trong com.dao"],
        ["Số entity", "21 file trong com.entity"],
    ])

    document.add_heading("2. Công nghệ sử dụng", level=1)
    add_bullets(document, [
        "Java 17 và Java Swing để xây dựng ứng dụng desktop.",
        "JDBC kết nối SQL Server thông qua mssql-jdbc.",
        "Maven quản lý build, dependency và đóng gói JAR.",
        "FlatLaf/FlatLaf Extras hỗ trợ giao diện Swing hiện đại.",
        "JCalendar hỗ trợ chọn ngày giờ trong các form nghiệp vụ.",
        "JFreeChart phục vụ thống kê và biểu đồ doanh thu.",
        "Apache POI phục vụ xuất dữ liệu Excel khi cần báo cáo.",
        "java-dotenv hỗ trợ cấu hình môi trường kết nối database.",
    ])

    document.add_heading("3. Kiến trúc tổng thể", level=1)
    document.add_paragraph(
        "Dự án được tổ chức theo hướng phân lớp rõ ràng: lớp giao diện trong modules, lớp truy cập dữ liệu trong dao, "
        "lớp mô hình dữ liệu trong entity, enum cho các trạng thái/loại dữ liệu cố định và connectDB phụ trách kết nối cơ sở dữ liệu."
    )
    add_table(document, ["Thành phần", "Vai trò"], [
        ["com.Main", "Điểm khởi chạy ứng dụng, thiết lập giao diện chính."],
        ["com.modules", "Các màn hình chức năng: bán vé, quản lý vé, hóa đơn, nhân viên, thống kê..."],
        ["com.dao", "Đóng gói truy vấn JDBC và thao tác CRUD với SQL Server."],
        ["com.entity", "Biểu diễn dữ liệu nghiệp vụ tương ứng các bảng database."],
        ["com.enums", "Khai báo vai trò, loại ghế, trạng thái vé, trạng thái nhân viên."],
        ["com.connectDB", "Quản lý kết nối tới SQL Server."],
        ["src/main/resources", "Chứa ảnh giao diện, tài liệu và script SQL khởi tạo dữ liệu."],
    ])

    document.add_heading("4. Chức năng chính", level=1)
    add_table(document, ["Chức năng", "Mô tả"], [
        ["Đăng nhập và phân quyền", "Hỗ trợ vai trò nhân viên bán vé, điều phối và quản trị theo enum VaiTro."],
        ["Bán vé nhiều bước", "Luồng BanVeModule và các BanVeStep*Module hỗ trợ chọn hành trình, lịch, ghế, xác nhận và thanh toán."],
        ["Quản lý vé", "Theo dõi trạng thái vé, thông tin khách hàng, tuyến, lịch và hỗ trợ hoàn/hủy vé."],
        ["Quản lý hóa đơn", "Lưu hóa đơn, chi tiết hóa đơn và giá tiền tại thời điểm bán."],
        ["Quản lý khách hàng/nhân viên", "Thêm, sửa, tìm kiếm, lọc và quản lý trạng thái dữ liệu người dùng."],
        ["Quản lý vận hành tàu", "Quản lý ga, tuyến, đầu máy, đoàn tàu, toa tàu, ghế và lịch chạy."],
        ["Quản lý giá và khuyến mãi", "Thiết lập kỳ giá, chi tiết giá, khuyến mãi và phạm vi áp dụng."],
        ["Thống kê", "Sử dụng JFreeChart để trình bày số liệu tổng quan/doanh thu."],
    ])

    document.add_heading("5. Cơ sở dữ liệu", level=1)
    document.add_paragraph(
        "Script SQL chính nằm tại src/main/resources/data/BanVeTauNhaGa_MSSQL.sql. "
        "Cơ sở dữ liệu gồm các bảng nghiệp vụ phục vụ bán vé, vận hành lịch tàu, hóa đơn, khách hàng, nhân viên, giá và khuyến mãi."
    )
    add_table(document, ["Nhóm bảng", "Ý nghĩa"], [
        ["NhanVien", "Thông tin nhân viên, vai trò và trạng thái làm việc."],
        ["KhachHang", "Thông tin khách hàng mua vé."],
        ["Ga, Tuyen", "Nhà ga và tuyến đường giữa ga đi/ga đến."],
        ["DauMay, DoanTau, ToaTau, Ghe", "Cấu hình đoàn tàu, toa và ghế/giường."],
        ["Lich", "Lịch chạy tàu, thời gian bắt đầu và thời gian chạy."],
        ["Gia, ChiTietGia", "Kỳ giá và giá niêm yết theo tuyến/loại ghế."],
        ["KhuyenMai, ChiTietKhuyenMai, ApDungKM", "Thông tin khuyến mãi và quan hệ áp dụng."],
        ["Ve, HoaDon, ChiTietHoaDon", "Vé, hóa đơn và chi tiết hóa đơn liên kết vé."],
        ["GiuCho", "Giữ chỗ tạm thời khi thao tác bán vé."],
    ])

    document.add_heading("6. Thiết kế giao diện", level=1)
    document.add_paragraph(
        "Giao diện dùng phong cách Azure Rail/Notion-like với design token tập trung trong NotionTheme và AppColors. "
        "Màu chủ đạo là xanh #005D90, nền sáng, viền nhẹ, typography rõ ràng và ưu tiên validation inline thay vì hộp thoại lỗi rời rạc."
    )
    add_bullets(document, [
        "Main frame chạy tối đa hóa, kích thước tối thiểu 1024x700.",
        "Các màn hình quản lý dùng bảng dữ liệu, thanh lọc/tìm kiếm và phân trang tự động.",
        "Dialog nhập liệu dùng lỗi inline, viền đỏ và tự động focus trường lỗi.",
        "Icon được vẽ bằng Java2D qua LineIcons, ảnh lớn đặt trong src/main/resources/images.",
        "Các module UI triển khai AppModule để dễ gắn vào menu/chuyển màn hình.",
    ])

    document.add_heading("7. Luồng bán vé đề xuất", level=1)
    steps = [
        "Nhân viên đăng nhập hệ thống.",
        "Chọn ga đi, ga đến, ngày đi và tìm lịch phù hợp.",
        "Chọn chuyến tàu, toa và ghế còn trống.",
        "Nhập hoặc chọn thông tin khách hàng.",
        "Tính giá theo tuyến, loại ghế, kỳ giá và khuyến mãi nếu có.",
        "Xác nhận hóa đơn và chọn phương thức thanh toán tiền mặt/chuyển khoản.",
        "Lưu vé, chi tiết hóa đơn và cập nhật trạng thái ghế/vé.",
    ]
    for index, step in enumerate(steps, start=1):
        document.add_paragraph(f"{index}. {step}")

    document.add_heading("8. Đánh giá hiện trạng", level=1)
    add_table(document, ["Điểm mạnh", "Ghi chú"], [
        ["Tổ chức code rõ tầng", "Tách modules, dao, entity, enums giúp dễ bảo trì."],
        ["Phạm vi nghiệp vụ rộng", "Bao phủ bán vé, hóa đơn, vận hành tàu, giá, khuyến mãi, thống kê."],
        ["Thiết kế UI thống nhất", "Có design token, module pattern và ghi chú kiến trúc UI."],
        ["Database tương đối đầy đủ", "Có các bảng chính và quan hệ phục vụ nghiệp vụ bán vé."],
    ])
    add_table(document, ["Rủi ro/Hạn chế", "Khuyến nghị"], [
        ["Phụ thuộc SQL Server local", "Cần chuẩn hóa file .env mẫu và hướng dẫn import database."],
        ["Chưa thấy test tự động", "Nên bổ sung test cho DAO/logic tính giá khi tách được khỏi UI."],
        ["Nhiều module UI lớn", "Nên tiếp tục tách component dùng chung để giảm trùng lặp."],
        ["Báo cáo Word này tạo từ khảo sát mã nguồn", "Cần bổ sung thông tin nhóm, yêu cầu môn học và ảnh màn hình nếu nộp chính thức."],
    ])

    document.add_heading("9. Hướng phát triển", level=1)
    add_bullets(document, [
        "Bổ sung tài liệu cài đặt chi tiết: JDK, Maven, SQL Server, biến môi trường và dữ liệu mẫu.",
        "Thêm ảnh màn hình cho các chức năng chính để báo cáo trực quan hơn.",
        "Tách nghiệp vụ tính giá/khuyến mãi thành service để dễ kiểm thử.",
        "Bổ sung phân quyền chi tiết theo vai trò và kiểm tra quyền tại từng module.",
        "Thêm chức năng xuất PDF/Excel cho hóa đơn, danh sách vé và thống kê doanh thu.",
        "Xây dựng bộ test hồi quy cho DAO, validation và luồng bán vé.",
    ])

    document.add_heading("10. Kết luận", level=1)
    document.add_paragraph(
        "Dự án BanVeTauNhaGa đã hình thành một ứng dụng desktop quản lý bán vé tàu tương đối hoàn chỉnh, "
        "có kiến trúc phân lớp, cơ sở dữ liệu rõ ràng và giao diện thống nhất. Với việc bổ sung tài liệu cài đặt, "
        "ảnh minh họa, kiểm thử tự động và tối ưu một số module lớn, dự án có thể tiếp tục phát triển thành sản phẩm ổn định hơn cho nghiệp vụ bán vé tại nhà ga."
    )

    add_page_number_footer(document)
    document.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
